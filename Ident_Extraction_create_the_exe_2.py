# -*- coding: utf-8 -*-
"""
Created on Tue Jan 23 06:56:28 2024

@author: KW

Title: Improved Ident_Extraction
"""

# -*- coding: utf-8 -*-
import os
from docx import Document
import re
import pandas as pd
import tkinter as tk
from tkinter import filedialog, messagebox

# Get data from word (function)
def get_data_from_word(path_to_file):
    doc_reader = Document(path_to_file)
    data = ""
    try:
        for p in doc_reader.paragraphs:
            data += p.text + "\n"
            for table in doc_reader.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            data += paragraph.text + "\n"
    except Exception as e:
        print(f"An error occurred: {e}")
    return data

# Reverse the Idents (function)
def custom_split(identifier):
    reversed_ident = identifier[::-1]
    parts = reversed_ident.split('-', maxsplit=3)

    if len(parts[0]) == 2:
        return parts[3][::-1], parts[2][::-1], parts[1][::-1], parts[0][::-1]
    elif len(parts[0]) == 3:
        return parts[2][::-1], parts[1][::-1], parts[0][::-1]
    else:
        return identifier.split('-')  # No pattern found, split normally

# Separate GUI-related functions
def get_file_path():
    root = tk.Tk()
    root.withdraw()
    return filedialog.askopenfilename(title="Select a Word file", filetypes=[("Word files", "*.docx")])

def save_file_dialog():
    return filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel files", "*.xlsx")], title="Save As")

def show_info(message):
    messagebox.showinfo("Success", message)

def show_warning(message):
    messagebox.showwarning("No Matches", message)

def show_error(message):
    messagebox.showerror("Error", message)

def process_data(extracted_data):
    pattern = r'\b(?:\w+-){2,}\w+\b'
    matches = re.findall(pattern, extracted_data)
    
    if not matches:
        show_warning("No matches found or pattern does not match the expected format.")
    else:
        unique_matches = pd.unique(matches).tolist()
        df = pd.DataFrame(unique_matches, columns=['ident'])
        df[['Dokument', 'Art', 'Teil', 'Vs']] = df['ident'].apply(lambda x: pd.Series(custom_split(x)))

        output_filename = save_file_dialog()

        if output_filename:
            df.to_excel(output_filename, index=False)
            show_info(f"Data exported to '{output_filename}'")
            
            # Open the created Excel file using the default application
            try:
                os.startfile(output_filename)
            except AttributeError:
                # For non-Windows systems or environments that don't support os.startfile()
                import subprocess
                subprocess.run(['xdg-open', output_filename], check=True)
        else:
            show_warning("Something went wrong.")

def process_file():
    file_path = get_file_path()
    if file_path:
        try:
            extracted_data = get_data_from_word(file_path)
            process_data(extracted_data)
        except Exception as e:
            show_error(f"An error occurred: {e}")
    else:
        show_warning("Please select a Word file.")

# Call the function to initiate the process
process_file()
